"""
生成需求文档模板 docx 文件。
运行方式：cd backend && uv run python ../scripts/create_requirement_template.py
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("python-docx 未安装，正在通过 uv 运行...")
    sys.exit(1)

OUTPUT_PATH = Path(__file__).parent.parent / "frontend" / "public" / "templates" / "requirement-template.docx"


def create_template() -> None:
    doc = Document()

    # Title
    title = doc.add_heading("需求标题（必填）", 0)
    title.style.font.color.rgb = RGBColor(0x1A, 0x20, 0x30)

    # Section: 功能描述
    doc.add_heading("功能描述", 1)
    doc.add_paragraph("（描述该需求的功能目标和业务背景）")

    # Section: 业务规则
    doc.add_heading("业务规则", 1)
    p = doc.add_paragraph()
    p.add_run("规则1：").bold = True
    p.add_run("描述业务规则内容")
    p2 = doc.add_paragraph()
    p2.add_run("规则2：").bold = True
    p2.add_run("描述业务规则内容")

    # Section: 输入条件
    doc.add_heading("输入条件", 1)
    doc.add_paragraph("（用户操作或系统触发条件，包括前置条件和触发事件）")

    # Section: 预期结果
    doc.add_heading("预期结果", 1)
    doc.add_paragraph("（系统应该如何响应，描述正常流程下的预期行为）")

    # Section: 异常场景
    doc.add_heading("异常场景", 1)
    doc.add_paragraph("（边界条件和错误处理，描述异常情况下的系统行为）")

    # Section: 非功能需求
    doc.add_heading("非功能需求", 1)
    doc.add_paragraph("性能：（如接口响应时间 ≤ 500ms）")
    doc.add_paragraph("安全：（如需要登录鉴权）")
    doc.add_paragraph("兼容性：（如支持 Chrome 90+）")

    # Section: 验收标准
    doc.add_heading("验收标准", 1)
    doc.add_paragraph("1. 验收条件1")
    doc.add_paragraph("2. 验收条件2")
    doc.add_paragraph("3. 验收条件3")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"已生成：{OUTPUT_PATH}")
    size = OUTPUT_PATH.stat().st_size
    print(f"文件大小：{size} bytes")


if __name__ == "__main__":
    create_template()
